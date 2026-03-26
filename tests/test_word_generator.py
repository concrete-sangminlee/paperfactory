import os
import sys
import json
import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from utils.word_generator import generate_word, load_guideline

MINIMAL_CONTENT = {
    "title": "Test Paper Title",
    "authors": "A. Author, B. Author",
    "abstract": "This is the abstract text for testing purposes.",
    "keywords": "keyword1; keyword2; keyword3",
    "sections": [
        {"heading": "INTRODUCTION", "content": "Introduction paragraph one.\n\nIntroduction paragraph two."},
        {"heading": "METHODOLOGY", "content": "Methodology text.", "subsections": [
            {"heading": "Data Collection", "content": "Sub-section text."},
        ]},
        {"heading": "CONCLUSIONS", "content": "Conclusions text."},
    ],
    "references": [
        "[1] A. Author, Title of paper, J. Eng. 50 (2024) 1-10. https://doi.org/10.1000/test.",
    ],
}


class TestGenerateWord:
    def test_creates_docx_file(self, tmp_path):
        content = {**MINIMAL_CONTENT}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        assert os.path.exists(path)
        assert path.endswith(".docx")

    def test_contains_title(self, tmp_path):
        content = {**MINIMAL_CONTENT}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Paper Title" in t for t in texts)

    def test_contains_abstract(self, tmp_path):
        content = {**MINIMAL_CONTENT}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("abstract text" in t.lower() for t in texts)

    def test_contains_references(self, tmp_path):
        content = {**MINIMAL_CONTENT}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("REFERENCES" in t for t in texts)

    def test_tables_inserted(self, tmp_path):
        content = {**MINIMAL_CONTENT, "tables": [
            {"caption": "Table 1. Test table.", "headers": ["Col A", "Col B"], "rows": [["1", "2"]]}
        ]}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        doc = Document(path)
        assert len(doc.tables) >= 1

    def test_elsevier_journal(self, tmp_path):
        content = {**MINIMAL_CONTENT}
        path = generate_word(content, "eng_structures", output_dir=str(tmp_path))
        assert os.path.exists(path)

    def test_data_availability_inserted_when_provided(self, tmp_path):
        content = {**MINIMAL_CONTENT, "data_availability": "Data available on request."}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Data available on request" in t for t in texts)

    def test_subsection_headings(self, tmp_path):
        content = {**MINIMAL_CONTENT}
        path = generate_word(content, "asce_jse", output_dir=str(tmp_path))
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Data Collection" in t for t in texts)


class TestLoadGuideline:
    def test_loads_asce(self):
        g = load_guideline("asce_jse")
        assert g["journal_name"] == "ASCE Journal of Structural Engineering"

    def test_loads_eng_structures(self):
        g = load_guideline("eng_structures")
        assert g["publisher"] == "Elsevier"
