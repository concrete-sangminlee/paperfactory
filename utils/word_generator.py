"""Word document generator for PaperFactory.

Supports all journals defined in guidelines/*.json.
"""

import json
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

GUIDELINES_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "guidelines"
)
DEFAULT_PAPERS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "outputs", "papers"
)


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------


def load_guideline(journal_key: str) -> dict:
    """Load and return the guideline JSON for *journal_key*."""
    path = os.path.join(GUIDELINES_DIR, f"{journal_key}.json")
    with open(path, "r") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _apply_font(run, formatting: dict) -> None:
    run.font.name = formatting.get("font", "Times New Roman")
    run.font.size = Pt(formatting.get("font_size", 12))


def _set_spacing(paragraph, formatting: dict) -> None:
    spacing = formatting.get("line_spacing", "double")
    pf = paragraph.paragraph_format
    if spacing == "double":
        pf.line_spacing = 2.0
    elif spacing in ("1.5", 1.5):
        pf.line_spacing = 1.5
    else:
        pf.line_spacing = 2.0


def _parse_margin(margin_str: str) -> float:
    """Return margin in inches parsed from a guideline margin string."""
    s = margin_str.lower()
    if "1 inch" in s or "2.54" in s or "1in" in s:
        return 1.0
    match = re.search(r"([\d.]+)\s*(inch|in|cm|mm)", s)
    if match:
        val = float(match.group(1))
        unit = match.group(2)
        if unit == "cm":
            return val / 2.54
        if unit == "mm":
            return val / 25.4
        return val
    return 1.0


def _add_heading_l1(doc, text: str, formatting: dict) -> None:
    """Level-1 heading: ALL CAPS, bold, left-aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text.upper())
    run.bold = True
    _apply_font(run, formatting)
    _set_spacing(para, formatting)


def _add_heading_l2(doc, text: str, formatting: dict) -> None:
    """Level-2 subsection heading: Title Case, bold, left-aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    _apply_font(run, formatting)
    _set_spacing(para, formatting)


def _add_heading_l3(doc, text: str, formatting: dict) -> None:
    """Level-3 subsection heading: Sentence case, italic, left-aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.italic = True
    _apply_font(run, formatting)
    _set_spacing(para, formatting)


def _add_body_text(doc, text: str, formatting: dict) -> None:
    """Add one or more body paragraphs, splitting on blank lines."""
    for chunk in text.split("\n\n"):
        chunk = chunk.strip()
        if not chunk:
            continue
        para = doc.add_paragraph()
        run = para.add_run(chunk)
        _apply_font(run, formatting)
        _set_spacing(para, formatting)


def _table_no_vertical_borders(table) -> None:
    """Remove vertical rules from *table* (horizontal-only style)."""
    tbl_xml = table._tbl
    tbl_pr = tbl_xml.tblPr if tbl_xml.tblPr is not None else tbl_xml._add_tblPr()
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "bottom", "insideH"):
        b = borders.find(qn(f"w:{name}"))
        if b is None:
            b = OxmlElement(f"w:{name}")
            borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
    for name in ("left", "right", "insideV"):
        b = borders.find(qn(f"w:{name}"))
        if b is None:
            b = OxmlElement(f"w:{name}")
            borders.append(b)
        b.set(qn("w:val"), "none")


def _table_full_borders(table) -> None:
    """Apply full grid borders to *table*."""
    tbl_xml = table._tbl
    tbl_pr = tbl_xml.tblPr if tbl_xml.tblPr is not None else tbl_xml._add_tblPr()
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "bottom", "left", "right", "insideH", "insideV"):
        b = borders.find(qn(f"w:{name}"))
        if b is None:
            b = OxmlElement(f"w:{name}")
            borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")


def _insert_table(doc, tbl_info: dict, formatting: dict, no_vertical: bool) -> None:
    """Render one table dict (caption above, data, optional border style)."""
    # Caption (above)
    cap_para = doc.add_paragraph()
    cap_run = cap_para.add_run(tbl_info.get("caption", ""))
    cap_run.bold = True
    _apply_font(cap_run, formatting)
    cap_run.font.size = Pt(formatting.get("font_size", 12) - 1)
    _set_spacing(cap_para, formatting)

    headers = tbl_info.get("headers", [])
    rows = tbl_info.get("rows", [])
    if not headers:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = formatting.get("font", "Times New Roman")
        run.font.size = Pt(formatting.get("font_size", 12) - 2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = formatting.get("font", "Times New Roman")
            run.font.size = Pt(formatting.get("font_size", 12) - 2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if no_vertical:
        _table_no_vertical_borders(table)
    else:
        _table_full_borders(table)

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Main public API
# ---------------------------------------------------------------------------


def generate_word(
    paper_content: dict,
    journal_key: str,
    figures: list = None,
    output_dir: str = None,
) -> str:
    """Generate a Word .docx file from *paper_content* following the journal guidelines.

    Parameters
    ----------
    paper_content:
        Dict with keys: title, authors, abstract, keywords, sections, tables,
        references, figure_captions, practical_applications, highlights,
        data_availability, acknowledgments, notation.
    journal_key:
        Key matching a file in guidelines/ (e.g. "asce_jse", "eng_structures").
    figures:
        Optional list of absolute paths to figure image files.
    output_dir:
        Directory in which to save the .docx file.  Defaults to
        outputs/papers/ inside the project root.

    Returns
    -------
    str
        Absolute path to the saved .docx file.
    """
    if output_dir is None:
        output_dir = DEFAULT_PAPERS_DIR
    os.makedirs(output_dir, exist_ok=True)

    guideline = load_guideline(journal_key)
    formatting = guideline.get("formatting", {})
    fig_config = guideline.get("figures_tables", {})
    publisher = guideline.get("publisher", "").lower()

    # Determine table border style from guideline
    table_rules = fig_config.get("table_rules", "")
    no_vertical_borders = (
        "no vertical" in table_rules.lower() or "no vertical rules" in table_rules.lower()
    )

    is_elsevier = "elsevier" in publisher
    is_asce = "american society of civil engineers" in publisher or "asce" in journal_key.lower()

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------
    doc = Document()

    section = doc.sections[0]
    page_size = formatting.get("page_size", "US Letter")
    if "A4" in page_size.upper():
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    margin_str = formatting.get("margins", "1 inch all sides")
    # "standard" is used by some guidelines — default to 1 inch
    if margin_str.lower() in ("standard", ""):
        margin_val = 1.0
    else:
        margin_val = _parse_margin(margin_str)
    section.top_margin = Inches(margin_val)
    section.bottom_margin = Inches(margin_val)
    section.left_margin = Inches(margin_val)
    section.right_margin = Inches(margin_val)

    # Global Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = formatting.get("font", "Times New Roman")
    normal_style.font.size = Pt(formatting.get("font_size", 12))

    # ------------------------------------------------------------------
    # Highlights (Elsevier-specific) — before title
    # ------------------------------------------------------------------
    highlights = paper_content.get("highlights", [])
    if is_elsevier and highlights:
        _add_heading_l1(doc, "Highlights", formatting)
        for bullet in highlights:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(str(bullet))
            _apply_font(run, formatting)
            _set_spacing(para, formatting)
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Title
    # ------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(paper_content.get("title", "Untitled"))
    title_run.bold = True
    title_run.font.size = Pt(formatting.get("font_size", 12) + 2)
    _apply_font(title_run, formatting)
    _set_spacing(title_para, formatting)

    # ------------------------------------------------------------------
    # Authors
    # ------------------------------------------------------------------
    if paper_content.get("authors"):
        authors_para = doc.add_paragraph()
        authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors_run = authors_para.add_run(paper_content["authors"])
        _apply_font(authors_run, formatting)
        _set_spacing(authors_para, formatting)

    # ------------------------------------------------------------------
    # Abstract
    # ------------------------------------------------------------------
    if paper_content.get("abstract"):
        doc.add_paragraph()
        _add_heading_l1(doc, "Abstract", formatting)
        _add_body_text(doc, paper_content["abstract"], formatting)

    # ------------------------------------------------------------------
    # Keywords
    # ------------------------------------------------------------------
    if paper_content.get("keywords"):
        kw_para = doc.add_paragraph()
        kw_label = kw_para.add_run("Keywords: ")
        kw_label.bold = True
        _apply_font(kw_label, formatting)
        kw_text = kw_para.add_run(paper_content["keywords"])
        _apply_font(kw_text, formatting)
        _set_spacing(kw_para, formatting)

    # ------------------------------------------------------------------
    # Practical Applications (ASCE-specific)
    # ------------------------------------------------------------------
    practical = paper_content.get("practical_applications", "")
    if is_asce and practical:
        doc.add_paragraph()
        _add_heading_l1(doc, "Practical Applications", formatting)
        _add_body_text(doc, practical, formatting)

    # ------------------------------------------------------------------
    # Body sections (with subsection support)
    # ------------------------------------------------------------------
    for sec in paper_content.get("sections", []):
        heading_text = sec.get("heading", "")
        body_text = sec.get("content", "")
        subsections = sec.get("subsections", [])

        doc.add_paragraph()
        _add_heading_l1(doc, heading_text, formatting)

        if body_text:
            _add_body_text(doc, body_text, formatting)

        for subsec in subsections:
            sub_heading = subsec.get("heading", "")
            sub_content = subsec.get("content", "")
            sub_subsections = subsec.get("subsections", [])

            _add_heading_l2(doc, sub_heading, formatting)
            if sub_content:
                _add_body_text(doc, sub_content, formatting)

            for subsubsec in sub_subsections:
                _add_heading_l3(doc, subsubsec.get("heading", ""), formatting)
                if subsubsec.get("content"):
                    _add_body_text(doc, subsubsec["content"], formatting)

    # ------------------------------------------------------------------
    # Notation section
    # ------------------------------------------------------------------
    notation = paper_content.get("notation", [])
    if notation:
        doc.add_paragraph()
        _add_heading_l1(doc, "Notation", formatting)
        # Build a two-column table: symbol | definition
        tbl = doc.add_table(rows=len(notation), cols=2)
        tbl.style = "Table Grid"
        for i, entry in enumerate(notation):
            symbol_cell = tbl.rows[i].cells[0]
            def_cell = tbl.rows[i].cells[1]
            for cell, text in (
                (symbol_cell, entry.get("symbol", "")),
                (def_cell, entry.get("definition", "")),
            ):
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(text)
                _apply_font(run, formatting)
        if no_vertical_borders:
            _table_no_vertical_borders(tbl)
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Tables
    # ------------------------------------------------------------------
    tables_data = paper_content.get("tables", [])
    if tables_data:
        doc.add_page_break()
        _add_heading_l1(doc, "Tables", formatting)
        for tbl_info in tables_data:
            _insert_table(doc, tbl_info, formatting, no_vertical_borders)

    # ------------------------------------------------------------------
    # Figures
    # ------------------------------------------------------------------
    placement = fig_config.get("placement", "end of manuscript")
    if figures and "end" in placement.lower():
        doc.add_page_break()
        _add_heading_l1(doc, "Figures", formatting)

    if figures:
        custom_captions = paper_content.get("figure_captions", [])
        cap_style = fig_config.get("figure_caption_style", "Fig. {n}. {desc}")
        for i, fig_path in enumerate(figures, 1):
            if os.path.exists(fig_path):
                fig_para = doc.add_paragraph()
                fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    fig_para.add_run().add_picture(fig_path, width=Inches(5.5))
                except Exception:
                    fig_para.add_run(f"[Figure {i}: {os.path.basename(fig_path)}]")
                _set_spacing(fig_para, formatting)

            if i <= len(custom_captions):
                cap_text = custom_captions[i - 1]
            else:
                desc = (
                    os.path.splitext(os.path.basename(fig_path))[0] if fig_path else f"Figure {i}"
                )
                cap_text = cap_style.replace("{n}", str(i)).replace("{desc}", desc)
            cap_para = doc.add_paragraph()
            cap_run = cap_para.add_run(cap_text)
            _apply_font(cap_run, formatting)
            _set_spacing(cap_para, formatting)

    # ------------------------------------------------------------------
    # Data Availability Statement
    # ------------------------------------------------------------------
    data_avail = paper_content.get("data_availability", "")
    if data_avail:
        doc.add_paragraph()
        _add_heading_l1(doc, "Data Availability Statement", formatting)
        _add_body_text(doc, data_avail, formatting)

    # ------------------------------------------------------------------
    # Acknowledgments
    # ------------------------------------------------------------------
    acknowledgments = paper_content.get("acknowledgments", "")
    if acknowledgments:
        doc.add_paragraph()
        _add_heading_l1(doc, "Acknowledgments", formatting)
        _add_body_text(doc, acknowledgments, formatting)

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------
    if paper_content.get("references"):
        doc.add_page_break()
        _add_heading_l1(doc, "References", formatting)

        refs = paper_content["references"]
        if isinstance(refs, str):
            refs = refs.split("\n")
        for ref in refs:
            ref = ref.strip()
            if ref:
                ref_para = doc.add_paragraph()
                ref_run = ref_para.add_run(ref)
                _apply_font(ref_run, formatting)
                _set_spacing(ref_para, formatting)

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    safe_title = re.sub(r"[^\w\s-]", "", paper_content.get("title", "paper"))[:50].strip()
    safe_title = re.sub(r"\s+", "_", safe_title)
    filename = f"{safe_title}_{journal_key}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path
